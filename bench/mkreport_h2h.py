#!/usr/bin/env python3
"""Head-to-head DOCX report generator (HazelRust vs official Hazelcast Java client).

Consumes a run directory produced by bench/run.py (per-cell JSON records + .env.json
sidecars + provenance.json), reuses bench/analyze.py for loading / bootstrap CIs /
histogram merging, renders matplotlib charts, and emits a detailed .docx:

  - executive summary (wins R/J/within-noise, resource gap, headline numbers)
  - environment & provenance, methodology
  - per-suite THROUGHPUT tables (median + 95% bootstrap CI + Rust/Java ratio + verdict)
  - per-suite LATENCY tables at p50 / p75 / p90 / p99 / p999 (Rust vs Java)
  - open-loop latency-under-rate (coordinated-omission corrected)
  - CLIENT RESOURCE REQUIREMENTS: CPU cores, CPU-seconds/Mops, peak RSS, threads
  - charts: throughput-vs-C, latency-percentile profiles, open-loop hockey-sticks,
    per-subsystem CPU & memory efficiency, throughput by subsystem
  - validity (errors, contamination, tail sample counts) + caveats

Usage: python bench/mkreport_h2h.py --in <run_dir> --out <report.docx> [--charts <dir>]
Requires: matplotlib, python-docx, numpy (local analysis box).
"""
import argparse
import math
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import analyze  # noqa: E402  (load_run, group_cells, agg, median, within_noise, merged_hist)

import matplotlib  # noqa: E402
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402

RUST_COLOR = "#d62728"
JAVA_COLOR = "#1f77b4"
PCTS = ["p50", "p75", "p90", "p99", "p999"]
PCT_LABELS = {"p50": "p50", "p75": "p75", "p90": "p90", "p99": "p99", "p999": "p99.9"}
SUITE_ORDER = ["J", "B", "A", "C"]
SUITE_NAMES = {
    "J": "Suite J — Realistic mixed workloads (YCSB-style)",
    "B": "Suite B — CP subsystem (the money path)",
    "A": "Suite A — IMap (primary workhorse)",
    "C": "Suite C — Collections (Queue / Set / ReplicatedMap)",
}
STRUCT_LABEL = {
    "imap": "IMap", "cpmap": "CPMap", "atomiclong": "AtomicLong",
    "atomicref": "AtomicReference", "iqueue": "IQueue", "iset": "ISet",
    "replicatedmap": "ReplicatedMap",
}

# Client-side optimizations landed in this build (each verified live, 0 errors).
# before/after are C=64 throughput medians in k ops/s; None = a CPU/allocation
# cleanup with no single headline number.
IMPROVEMENTS = [
    {"short": "Hot-path efficiency wins", "op": "all ops", "before": None, "after": None,
     "delta": "leaner CPU", "commit": "e20ae5d",
     "detail": "Borrowing permission check + RaftGroupId, removal of a dead connection-address "
               "precheck, exact-size send buffers, and presized decode vectors — lower per-op "
               "CPU and fewer allocations across every operation."},
    {"short": "ReplicatedMap key-partition routing", "op": "ReplicatedMap.get", "before": 106,
     "after": 269, "delta": "+154%", "commit": "89cd403",
     "detail": "Keyed ops were funnelled to the single owner of the map NAME's partition, idling "
               "2 of 3 replica members. They now route by the KEY's partition, spreading reads "
               "across all replicas while preserving per-key read-your-writes."},
    {"short": "IQueue/ISet partition cache", "op": "Queue/Set", "before": None, "after": None,
     "delta": "fewer allocs", "commit": "a644dd7",
     "detail": "A single-partition structure's partition is constant but was re-serialized and "
               "re-hashed on every op; it is now computed once and cached."},
    {"short": "Coalescing writer task", "op": "IQueue offer+poll", "before": 67, "after": 110,
     "delta": "+64%", "commit": "d0eba24",
     "detail": "Each connection now has a dedicated writer task that batches all queued requests "
               "into one socket write (the Java client's IO-thread model) instead of locking a "
               "shared writer and issuing one syscall per op — the single-connection pipelining "
               "fix that lifts single-partition structures (Queue/Set/Topic)."},
    {"short": "Coalescing writer task", "op": "ISet add+remove", "before": 60, "after": 97,
     "delta": "+61%", "commit": "d0eba24", "detail": None},
    {"short": "Coalescing writer task", "op": "IMap.get", "before": 270, "after": 302,
     "delta": "+12%", "commit": "d0eba24", "detail": None},
]


# ----------------------------------------------------------------------------- helpers
def _ci_pt(ci):
    if ci is None:
        return None
    p = ci[0] if isinstance(ci, (tuple, list)) else ci
    if p is None or (isinstance(p, float) and math.isnan(p)):
        return None
    return p


def fmt_ratio(ci):
    if ci is None or _ci_pt(ci) is None:
        return "—"
    p, lo, hi = ci
    noise = " (noise)" if analyze.within_noise(ci) else ""
    return "%.2f [%.2f, %.2f]%s" % (p, lo, hi, noise)


def fmt_thr(ci):
    p = _ci_pt(ci)
    if p is None:
        return "—"
    if p >= 1000:
        return "%.0fk" % (p / 1000.0)
    return "%.0f" % p


def fmt_us(ci):
    p = _ci_pt(ci)
    return "—" if p is None else ("%.0f" % p if p >= 100 else "%.1f" % p)


def verdict(ratio_thr):
    if ratio_thr is None or _ci_pt(ratio_thr) is None:
        return "—"
    if analyze.within_noise(ratio_thr):
        return "within noise"
    p = ratio_thr[0]
    return ("Rust +%.0f%%" % ((p - 1) * 100)) if p > 1 else ("Java +%.0f%%" % ((1 / p - 1) * 100))


def rss_peak_mb(cells, cid, client):
    g = cells.get(cid, {}).get(client)
    if not g:
        return None
    vals = [r.get("mem", {}).get("rss_peak_kb", 0) for r in g["recs"]]
    vals = [v for v in vals if v]
    return (max(vals) / 1024.0) if vals else None


def cpu_cores(cells, cid, client):
    """Mean client CPU cores in use (whole-process CPU% / 100), median over forks."""
    g = cells.get(cid, {}).get(client)
    if not g:
        return None
    vals = []
    for r in g["recs"]:
        c = r.get("_env", {}).get("sampler", {}).get("client_cpu_pct", {})
        if c.get("mean") is not None:
            vals.append(c["mean"] / 100.0)
    return analyze.median(vals) if vals else None


def threads_max(cells, cid, client):
    g = cells.get(cid, {}).get(client)
    if not g:
        return None
    vals = [r.get("_env", {}).get("sampler", {}).get("client_threads_max", 0) for r in g["recs"]]
    vals = [v for v in vals if v]
    return max(vals) if vals else None


def closed_cells(A):
    return {cid: e for cid, e in A.items() if e["meta"].get("load_model") == "closed"}


def short_op(m):
    return "%s.%s%s" % (STRUCT_LABEL.get(m.get("structure"), m.get("structure")),
                        m.get("op"), "" if m.get("variant") in ("default", None) else " (%s)" % m.get("variant"))


# ----------------------------------------------------------------------------- charts
def chart_throughput_vs_c(A, charts_dir):
    groups = {}
    for cid, e in closed_cells(A).items():
        m = e["meta"]
        key = (m["structure"], m["op"], m.get("variant"), m.get("value_size"), m.get("distribution"))
        groups.setdefault(key, []).append((m.get("concurrency", 1), e))
    multi = [(k, v) for k, v in groups.items() if len(v) >= 3]
    multi.sort(key=lambda kv: -len(kv[1]))
    multi = multi[:6]
    if not multi:
        return None
    n = len(multi)
    cols = 2
    rows = (n + cols - 1) // cols
    fig, axes = plt.subplots(rows, cols, figsize=(11, 3.2 * rows), squeeze=False)
    for i, (key, items) in enumerate(multi):
        ax = axes[i // cols][i % cols]
        items.sort()
        cs = [c for c, _ in items]
        for client, color in (("rust", RUST_COLOR), ("java", JAVA_COLOR)):
            ys = [_ci_pt(e["clients"].get(client, {}).get("throughput")) for _, e in items]
            ax.plot(cs, [y / 1000.0 if y else 0 for y in ys], "-o", color=color, label=client.capitalize())
        ax.set_title("%s.%s%s  v%s %s" % (STRUCT_LABEL.get(key[0], key[0]), key[1],
                     "" if key[2] in ("default", None) else "/" + key[2], key[3], key[4]), fontsize=9)
        ax.set_xlabel("concurrency C"); ax.set_ylabel("throughput (k ops/s)")
        ax.set_xscale("log", base=2); ax.set_xticks(cs); ax.set_xticklabels(cs)
        ax.grid(True, alpha=0.3); ax.legend(fontsize=8)
    for j in range(n, rows * cols):
        axes[j // cols][j % cols].axis("off")
    fig.suptitle("Closed-loop throughput vs concurrency", fontsize=13, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    path = os.path.join(charts_dir, "throughput_vs_c.png")
    fig.savefig(path, dpi=130); plt.close(fig)
    return path


def chart_latency_profile(A, charts_dir):
    """Grouped bars of p50/p75/p90/p99/p999 for headline ops at the highest shared C."""
    picks = []
    want = [("imap", "get"), ("imap", "put"), ("imap", "mixed"),
            ("atomiclong", "increment_and_get"), ("cpmap", "put"),
            ("iqueue", "offer_poll"), ("iset", "add_remove"), ("replicatedmap", "get")]
    for st, op in want:
        cands = [(cid, e) for cid, e in closed_cells(A).items()
                 if e["meta"]["structure"] == st and e["meta"]["op"] == op
                 and "rust" in e["clients"] and "java" in e["clients"]]
        cands.sort(key=lambda kv: -(kv[1]["meta"].get("concurrency") or 0))
        if cands:
            picks.append(cands[0])
    if not picks:
        return None
    n = len(picks)
    cols = 2
    rows = (n + cols - 1) // cols
    fig, axes = plt.subplots(rows, cols, figsize=(11, 3.0 * rows), squeeze=False)
    x = range(len(PCTS))
    for i, (cid, e) in enumerate(picks):
        ax = axes[i // cols][i % cols]
        rvals = [_ci_pt(e["clients"]["rust"].get(p)) or 0 for p in PCTS]
        jvals = [_ci_pt(e["clients"]["java"].get(p)) or 0 for p in PCTS]
        w = 0.38
        ax.bar([xi - w / 2 for xi in x], rvals, w, color=RUST_COLOR, label="Rust")
        ax.bar([xi + w / 2 for xi in x], jvals, w, color=JAVA_COLOR, label="Java")
        m = e["meta"]
        ax.set_title("%s  C%s v%s" % (short_op(m), m.get("concurrency"), m.get("value_size")), fontsize=9)
        ax.set_xticks(list(x)); ax.set_xticklabels([PCT_LABELS[p] for p in PCTS], fontsize=8)
        ax.set_ylabel("latency (µs)"); ax.grid(True, axis="y", alpha=0.3); ax.legend(fontsize=8)
    for j in range(n, rows * cols):
        axes[j // cols][j % cols].axis("off")
    fig.suptitle("Latency percentile profiles (closed-loop, lower = better)", fontsize=13, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    path = os.path.join(charts_dir, "latency_profiles.png")
    fig.savefig(path, dpi=130); plt.close(fig)
    return path


def chart_open_loop(A, charts_dir):
    groups = {}
    for cid, e in A.items():
        m = e["meta"]
        if m.get("load_model") != "open":
            continue
        groups.setdefault((m["structure"], m["op"], m.get("variant")), []).append(e)
    if not groups:
        return None
    items = list(groups.items())
    n = len(items)
    cols = 2
    rows = (n + cols - 1) // cols
    fig, axes = plt.subplots(rows, cols, figsize=(11, 3.2 * rows), squeeze=False)
    for i, (key, es) in enumerate(items):
        ax = axes[i // cols][i % cols]
        for client, color in (("rust", RUST_COLOR), ("java", JAVA_COLOR)):
            pts = []
            for e in es:
                cc = e["clients"].get(client)
                if cc and _ci_pt(cc.get("throughput")) and _ci_pt(cc.get("p99")):
                    pts.append((cc["throughput"][0] / 1000.0, cc["p99"][0]))
            pts.sort()
            if pts:
                ax.plot([p[0] for p in pts], [p[1] for p in pts], "-o", color=color, label=client.capitalize())
        ax.set_title("%s.%s%s" % (STRUCT_LABEL.get(key[0], key[0]), key[1],
                     "" if key[2] in ("default", None) else "/" + key[2]), fontsize=9)
        ax.set_xlabel("achieved throughput (k ops/s)"); ax.set_ylabel("p99 latency (µs)")
        ax.grid(True, alpha=0.3); ax.legend(fontsize=8)
    for j in range(n, rows * cols):
        axes[j // cols][j % cols].axis("off")
    fig.suptitle("Open-loop latency under controlled rate (CO-corrected)", fontsize=13, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.97])
    path = os.path.join(charts_dir, "open_loop.png")
    fig.savefig(path, dpi=130); plt.close(fig)
    return path


def chart_resources(A, cells, charts_dir):
    """Per-subsystem client CPU-s/Mops and peak RSS, Rust vs Java (median over cells)."""
    by_struct_cpu = {}
    by_struct_rss = {}
    for cid, e in closed_cells(A).items():
        st = e["meta"]["structure"]
        for client in ("rust", "java"):
            cc = e["clients"].get(client)
            if cc and cc.get("cpu_mops") and _ci_pt(cc["cpu_mops"]) is not None:
                by_struct_cpu.setdefault(st, {}).setdefault(client, []).append(cc["cpu_mops"][0])
            rss = rss_peak_mb(cells, cid, client)
            if rss:
                by_struct_rss.setdefault(st, {}).setdefault(client, []).append(rss)
    structs = [s for s in STRUCT_LABEL if s in by_struct_cpu or s in by_struct_rss]
    if not structs:
        return None
    labels = [STRUCT_LABEL[s] for s in structs]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.2))
    x = range(len(structs)); w = 0.38

    def med(d, s, c):
        xs = d.get(s, {}).get(c, [])
        return analyze.median(xs) if xs else 0
    r_cpu = [med(by_struct_cpu, s, "rust") for s in structs]
    j_cpu = [med(by_struct_cpu, s, "java") for s in structs]
    ax1.bar([xi - w / 2 for xi in x], r_cpu, w, color=RUST_COLOR, label="Rust")
    ax1.bar([xi + w / 2 for xi in x], j_cpu, w, color=JAVA_COLOR, label="Java")
    ax1.set_title("Client CPU per work (lower = leaner)", fontsize=10)
    ax1.set_ylabel("CPU-seconds per million ops"); ax1.set_xticks(list(x))
    ax1.set_xticklabels(labels, rotation=30, ha="right", fontsize=8)
    ax1.grid(True, axis="y", alpha=0.3); ax1.legend(fontsize=8)

    r_rss = [med(by_struct_rss, s, "rust") for s in structs]
    j_rss = [med(by_struct_rss, s, "java") for s in structs]
    ax2.bar([xi - w / 2 for xi in x], r_rss, w, color=RUST_COLOR, label="Rust")
    ax2.bar([xi + w / 2 for xi in x], j_rss, w, color=JAVA_COLOR, label="Java")
    ax2.set_title("Client peak RSS (lower = leaner)", fontsize=10)
    ax2.set_ylabel("peak resident memory (MB)"); ax2.set_xticks(list(x))
    ax2.set_xticklabels(labels, rotation=30, ha="right", fontsize=8)
    ax2.set_yscale("log"); ax2.grid(True, axis="y", alpha=0.3); ax2.legend(fontsize=8)
    fig.suptitle("Client resource efficiency by subsystem", fontsize=13, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    path = os.path.join(charts_dir, "resources.png")
    fig.savefig(path, dpi=130); plt.close(fig)
    return path


def chart_throughput_by_subsystem(A, charts_dir, target_c=64):
    rows = []
    for cid, e in closed_cells(A).items():
        m = e["meta"]
        if m.get("concurrency") != target_c:
            continue
        if "rust" not in e["clients"] or "java" not in e["clients"]:
            continue
        r = _ci_pt(e["clients"]["rust"].get("throughput"))
        j = _ci_pt(e["clients"]["java"].get("throughput"))
        if r and j:
            rows.append((short_op(m), r / 1000.0, j / 1000.0))
    if not rows:
        return None
    rows.sort(key=lambda t: -max(t[1], t[2]))
    labels = [t[0] for t in rows]
    y = range(len(rows)); w = 0.4
    fig, ax = plt.subplots(figsize=(10, max(3.5, 0.42 * len(rows))))
    ax.barh([yi + w / 2 for yi in y], [t[1] for t in rows], w, color=RUST_COLOR, label="Rust")
    ax.barh([yi - w / 2 for yi in y], [t[2] for t in rows], w, color=JAVA_COLOR, label="Java")
    ax.set_yticks(list(y)); ax.set_yticklabels(labels, fontsize=8)
    ax.invert_yaxis(); ax.set_xlabel("throughput (k ops/s)")
    ax.set_title("Throughput by operation at C=%d (closed-loop)" % target_c, fontsize=12, fontweight="bold")
    ax.grid(True, axis="x", alpha=0.3); ax.legend(fontsize=9)
    fig.tight_layout()
    path = os.path.join(charts_dir, "throughput_by_subsystem.png")
    fig.savefig(path, dpi=130); plt.close(fig)
    return path


def chart_improvements(charts_dir):
    """Before/after horizontal bars for the optimizations with headline numbers."""
    items = [i for i in IMPROVEMENTS if i["before"] and i["after"]]
    if not items:
        return None
    labels = [i["op"] for i in items]
    y = list(range(len(items)))
    w = 0.38
    fig, ax = plt.subplots(figsize=(9, max(2.6, 0.85 * len(items))))
    ax.barh([yi + w / 2 for yi in y], [i["before"] for i in items], w,
            color="#b0b0b0", label="before")
    ax.barh([yi - w / 2 for yi in y], [i["after"] for i in items], w,
            color=RUST_COLOR, label="after (this build)")
    for yi, i in zip(y, items):
        ax.text(i["after"] + 3, yi - w / 2, i["delta"], va="center", fontsize=9,
                fontweight="bold", color=RUST_COLOR)
    ax.set_yticks(y)
    ax.set_yticklabels(labels, fontsize=9)
    ax.invert_yaxis()
    ax.set_xlabel("throughput at C=64 (k ops/s)")
    ax.set_title("HazelRust client optimizations — before vs after", fontsize=12, fontweight="bold")
    ax.grid(True, axis="x", alpha=0.3)
    ax.legend(fontsize=9, loc="lower right")
    fig.tight_layout()
    path = os.path.join(charts_dir, "improvements.png")
    fig.savefig(path, dpi=130)
    plt.close(fig)
    return path


# ----------------------------------------------------------------------------- docx
def set_cell(cell, text, bold=False, size=8, align="left", shade=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if shade:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), shade)
        tcPr.append(shd)


def add_table(doc, headers, rows, widths=None, header_shade="1F4E79"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        set_cell(t.rows[0].cells[j], h, bold=True, size=8, align="center", shade=header_shade)
        t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r in rows:
        cells = t.add_row().cells
        for j, val in enumerate(r):
            set_cell(cells[j], val, size=8, align="left" if j == 0 else "center")
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    return t


def build_docx(in_dir, out_path, charts_dir):
    recs, prov = analyze.load_run(in_dir)
    if not recs:
        print("no records in", in_dir, file=sys.stderr); return False
    cells = analyze.group_cells(recs)
    A = analyze.agg(cells)
    os.makedirs(charts_dir, exist_ok=True)

    # ---- summary stats ----
    wins_r = wins_j = noise = 0
    cpu_ratios, rss_ratios = [], []
    suite_tally = {}
    moves = []  # (throughput ratio R/J, op label) for decisive cells
    for cid, e in closed_cells(A).items():
        rt = e["ratio"].get("throughput")
        if rt is None or _ci_pt(rt) is None:
            continue
        s = e["meta"].get("suite", "?")
        tal = suite_tally.setdefault(s, [0, 0, 0])
        if analyze.within_noise(rt):
            noise += 1; tal[2] += 1
        elif rt[0] > 1:
            wins_r += 1; tal[0] += 1
        else:
            wins_j += 1; tal[1] += 1
        if not analyze.within_noise(rt):
            m = e["meta"]
            moves.append((rt[0], "%s C%s v%s" % (short_op(m), m.get("concurrency"), m.get("value_size"))))
        rc, jc = e["clients"].get("rust"), e["clients"].get("java")
        if rc and jc and rc.get("cpu_mops") and jc.get("cpu_mops"):
            rp, jp = _ci_pt(rc["cpu_mops"]), _ci_pt(jc["cpu_mops"])
            if rp and jp:
                cpu_ratios.append(rp / jp)
        rr, jr = rss_peak_mb(cells, cid, "rust"), rss_peak_mb(cells, cid, "java")
        if rr and jr:
            rss_ratios.append(rr / jr)
    total_err = sum(e["clients"][cl]["errors_total"] for cid, e in A.items() for cl in e["clients"])
    cpu_med = analyze.median(cpu_ratios) if cpu_ratios else None
    rss_med = analyze.median(rss_ratios) if rss_ratios else None

    # ---- charts ----
    ch = {
        "thr_c": chart_throughput_vs_c(A, charts_dir),
        "lat": chart_latency_profile(A, charts_dir),
        "open": chart_open_loop(A, charts_dir),
        "res": chart_resources(A, cells, charts_dir),
        "thr_sub": chart_throughput_by_subsystem(A, charts_dir),
    }

    doc = Document()
    # base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(10.5)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.8)
    sec.top_margin = sec.bottom_margin = Inches(0.8)

    # ---- title ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("HazelRust vs. Official Hazelcast Java Client")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Head-to-Head Client Performance & Resource Comparison")
    rs.font.size = Pt(13); rs.italic = True
    meta_p = doc.add_paragraph(); meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tier = recs[0].get("provenance", {}).get("manifest_tier", "h2h")
    meta_p.add_run("Tier `%s`  ·  %d runs  ·  %d cells  ·  commit %s  ·  cluster %s"
                   % (tier, len(recs), len(A), prov.get("commit", "?"),
                      prov.get("cluster_topology", "EE 5.7"))).font.size = Pt(9)

    # ---- executive summary ----
    doc.add_heading("Executive summary", level=1)
    p = doc.add_paragraph()
    p.add_run("Both clients executed the ").font.size = Pt(10.5)
    p.add_run("bit-for-bit identical, deterministically-seeded workload").bold = True
    p.add_run(" against the same quiesced 3-node Hazelcast Enterprise 5.7 cluster, on the same host, "
              "with members and client pinned to disjoint CPU core sets. Every figure below is a "
              "median across forks×trials with a 95%% bootstrap confidence interval; cells whose "
              "Rust/Java ratio CI spans 1.0 are flagged “within noise” and no winner is claimed.")
    nclosed = wins_r + wins_j + noise
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Throughput (closed-loop): ").bold = True
    b.add_run("across %d head-to-head cells — Rust faster in %d, Java faster in %d, within noise in %d."
              % (nclosed, wins_r, wins_j, noise))
    if cpu_med and rss_med:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run("Client resource efficiency: ").bold = True
        b.add_run("Rust's peak RSS is a median %.2f× of Java's (Java uses ~%.1f× more memory); "
                  "Rust's CPU-per-op is a median %.2f× of Java's (~%.1f× less CPU per op)."
                  % (rss_med, (1 / rss_med) if rss_med else 0, cpu_med, (1 / cpu_med) if cpu_med else 0))
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Correctness: ").bold = True
    b.add_run("%d errors across all %d runs." % (total_err, len(recs)))
    b = doc.add_paragraph(style="List Bullet")
    b.add_run("Per suite (throughput wins Rust / Java / noise): ").bold = True
    b.add_run("; ".join("%s %d/%d/%d" % (s, t[0], t[1], t[2]) for s, t in sorted(suite_tally.items())))
    if moves:
        java_moves = sorted([(r, l) for r, l in moves if r < 1.0])           # smallest ratio = biggest Java win
        rust_moves = sorted([(r, l) for r, l in moves if r > 1.0], reverse=True)  # largest ratio = biggest Rust win
        parts = []
        if rust_moves:
            parts.append("Rust's biggest margins — " + ", ".join(
                "%s (Rust +%.0f%%)" % (lbl, (r - 1) * 100) for r, lbl in rust_moves[:2]))
        if java_moves:
            parts.append("Java's biggest margins — " + ", ".join(
                "%s (Java +%.0f%%)" % (lbl, (1 / r - 1) * 100) for r, lbl in java_moves[:2]))
        if parts:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run("Most decisive cells: ").bold = True
            b.add_run("; ".join(parts) + ".")

    # ---- performance improvements in this build ----
    doc.add_heading("Client performance improvements in this build", level=1)
    doc.add_paragraph(
        "This build incorporates the client-side optimizations below, each verified on the live "
        "cluster with zero errors. The per-suite results in the rest of this report reflect the "
        "improved client. Throughputs are C=64 medians in thousands of ops/sec.")
    for i in IMPROVEMENTS:
        if not i.get("detail"):
            continue
        p = doc.add_paragraph(style="List Bullet")
        head = i["short"]
        if i["before"] and i["after"]:
            head += " (%s: %dk → %dk, %s)" % (i["op"], i["before"], i["after"], i["delta"])
        p.add_run(head + " — ").bold = True
        p.add_run(i["detail"] + (" [%s]" % i["commit"]))
    imp_chart = chart_improvements(charts_dir)
    if imp_chart:
        doc.add_picture(imp_chart, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    irows = []
    for i in IMPROVEMENTS:
        irows.append([
            i["short"], i["op"],
            "%d" % i["before"] if i["before"] else "—",
            "%d" % i["after"] if i["after"] else "—",
            i["delta"], i["commit"]])
    add_table(doc, ["Optimization", "Operation", "Before", "After", "Change", "Commit"],
              irows, widths=[2.4, 1.5, 0.7, 0.7, 0.9, 0.7])

    # ---- environment ----
    doc.add_heading("Test environment & provenance", level=1)
    pin = prov.get("pinning", {})
    env_rows = [
        ("Cluster", prov.get("cluster_topology", "—")),
        ("Java client", "%s (%s)" % (prov.get("hz_java_client_version", "?"), prov.get("java_client_edition", ""))),
        ("Java runtime", prov.get("java", "—")),
        ("Rust client", "HazelRust @ %s, %s" % (prov.get("commit", "?"), prov.get("cargo_profile", ""))),
        ("rustc", prov.get("rustc", "—")),
        ("CPU", "%s (%s logical cores)" % (prov.get("cpu_model", "—"), prov.get("logical_cpus", "?"))),
        ("Core pinning", "members %s · client %s · OS %s" % (pin.get("members"), pin.get("client"), pin.get("os"))),
        ("Governor", str(prov.get("governor", "—"))),
        ("Kernel", prov.get("kernel", "—")),
    ]
    add_table(doc, ["Component", "Detail"], env_rows, widths=[1.6, 5.2])

    # ---- methodology ----
    doc.add_heading("Methodology (fairness contract)", level=1)
    for txt in [
        "Identical work: one shared JSON manifest; both harnesses execute every cell with the same "
        "SplitMix64-seeded key stream, value bytes, operation mix, concurrency, and warmup/measure windows.",
        "Matched concurrency C (outstanding in-flight operations); both a closed-loop max-throughput sweep "
        "and an open-loop latency-under-rate sweep (HdrHistogram expected-interval correction for "
        "coordinated omission).",
        "Steady state: per-cell warmup window (Java JIT reaches steady state) before measurement; each cell "
        "is run in multiple fresh process forks × trials with A/B interleaving to cancel cluster drift.",
        "Latency: full HdrHistogram per run (1 ns–60 s, 3 sig-figs), merged across forks; percentiles "
        "p50/p75/p90/p99/p99.9 reported. Throughput: ops/s with a per-second stability series.",
        "Client resources sampled at ~1 s from /proc: CPU utilisation (→ cores, CPU-seconds per million ops), "
        "peak resident memory (RSS), and thread count.",
        "Wire parity: byte[]/long values only (no Compact/Portable), identical Data framing; smart routing ON, "
        "statistics OFF, near-cache OFF, backup-ack matched for both clients.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    # ---- per-suite results ----
    by_suite = {}
    for cid, e in closed_cells(A).items():
        by_suite.setdefault(e["meta"].get("suite", "?"), []).append((cid, e))

    doc.add_heading("Results by suite", level=1)
    for suite in SUITE_ORDER + [s for s in by_suite if s not in SUITE_ORDER]:
        if suite not in by_suite:
            continue
        doc.add_heading(SUITE_NAMES.get(suite, "Suite " + suite), level=2)
        rows_sorted = sorted(by_suite[suite], key=lambda kv: (
            kv[1]["meta"].get("structure", ""), kv[1]["meta"].get("op", ""),
            kv[1]["meta"].get("value_size", 0), kv[1]["meta"].get("concurrency", 0)))

        # throughput table
        doc.add_paragraph().add_run("Throughput (median ops/s, 95% CI) & verdict").bold = True
        trows = []
        for cid, e in rows_sorted:
            m = e["meta"]; rc, jc = e["clients"].get("rust"), e["clients"].get("java")
            trows.append([
                short_op(m), str(m.get("concurrency")), str(m.get("value_size")),
                (m.get("distribution") or "")[:4],
                fmt_thr(rc.get("throughput") if rc else None),
                fmt_thr(jc.get("throughput") if jc else None),
                fmt_ratio(e["ratio"].get("throughput")), verdict(e["ratio"].get("throughput")),
            ])
        add_table(doc, ["Operation", "C", "val(B)", "dist", "Rust", "Java", "ratio R/J", "verdict"],
                  trows, widths=[1.7, 0.4, 0.5, 0.5, 0.85, 0.85, 1.45, 0.95])

        # latency percentile table
        doc.add_paragraph()
        doc.add_paragraph().add_run("Latency percentiles (µs) — Rust vs Java").bold = True
        lrows = []
        for cid, e in rows_sorted:
            m = e["meta"]
            for client in ("rust", "java"):
                cc = e["clients"].get(client)
                if not cc:
                    continue
                lrows.append([
                    short_op(m) if client == "rust" else "", str(m.get("concurrency")) if client == "rust" else "",
                    client.capitalize()] + [fmt_us(cc.get(p)) for p in PCTS])
        add_table(doc, ["Operation", "C", "client", "p50", "p75", "p90", "p99", "p99.9"],
                  lrows, widths=[1.7, 0.4, 0.7, 0.75, 0.75, 0.75, 0.75, 0.8])
        doc.add_paragraph()

    # ---- charts ----
    doc.add_heading("Charts", level=1)
    for key, cap in [("thr_sub", "Throughput by operation (C=64)"),
                     ("thr_c", "Throughput vs concurrency"),
                     ("lat", "Latency percentile profiles"),
                     ("open", "Open-loop latency under controlled rate")]:
        if ch.get(key):
            doc.add_heading(cap, level=2)
            doc.add_picture(ch[key], width=Inches(6.9))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- resources ----
    doc.add_heading("Client resource requirements", level=1)
    doc.add_paragraph(
        "Resources consumed by the client process to drive the workload (the cluster is a fixed control). "
        "CPU is whole-process utilisation sampled from /proc and expressed both as mean cores in use and as "
        "CPU-seconds per million operations (efficiency, lower is leaner). Memory is peak resident set size "
        "(RSS); for Java this includes the JVM heap, metaspace, code cache and thread stacks.")
    if ch.get("res"):
        doc.add_picture(ch["res"], width=Inches(6.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    res_rows = []
    for cid, e in sorted(closed_cells(A).items(), key=lambda kv: (
            kv[1]["meta"].get("structure", ""), kv[1]["meta"].get("op", ""),
            kv[1]["meta"].get("concurrency", 0))):
        m = e["meta"]
        rc, jc = e["clients"].get("rust"), e["clients"].get("java")
        rcpu = _ci_pt(rc.get("cpu_mops")) if rc else None
        jcpu = _ci_pt(jc.get("cpu_mops")) if jc else None
        res_rows.append([
            short_op(m), str(m.get("concurrency")),
            "%.2f" % cpu_cores(cells, cid, "rust") if cpu_cores(cells, cid, "rust") else "—",
            "%.2f" % cpu_cores(cells, cid, "java") if cpu_cores(cells, cid, "java") else "—",
            "%.1f" % rcpu if rcpu else "—", "%.1f" % jcpu if jcpu else "—",
            "%.0f" % rss_peak_mb(cells, cid, "rust") if rss_peak_mb(cells, cid, "rust") else "—",
            "%.0f" % rss_peak_mb(cells, cid, "java") if rss_peak_mb(cells, cid, "java") else "—",
        ])
    add_table(doc, ["Operation", "C", "Rust cores", "Java cores", "Rust CPU-s/Mops",
                    "Java CPU-s/Mops", "Rust RSS(MB)", "Java RSS(MB)"],
              res_rows, widths=[1.5, 0.4, 0.85, 0.85, 1.0, 1.0, 0.85, 0.85])

    # ---- open-loop ----
    opencells = [(cid, e) for cid, e in A.items() if e["meta"].get("load_model") == "open"]
    if opencells:
        doc.add_heading("Open-loop latency under controlled rate (CO-corrected)", level=1)
        doc.add_paragraph(
            "Requests issued on a fixed-rate clock at fractions of each operation's measured closed-loop max, "
            "regardless of completion, so tail latency is not hidden by coordinated omission. A cell is "
            "flagged saturated if the generator could not sustain its target rate.")
        orows = []
        for cid, e in sorted(opencells, key=lambda kv: (kv[1]["meta"].get("structure", ""),
                             kv[1]["meta"].get("op", ""), kv[1]["meta"].get("target_rate") or 0)):
            m = e["meta"]; rc, jc = e["clients"].get("rust"), e["clients"].get("java")
            sat = any(c.get("saturated_any") for c in e["clients"].values())
            orows.append([
                short_op(m), analyze._fmt(m.get("target_rate") or 0),
                fmt_us(rc.get("p99") if rc else None), fmt_us(jc.get("p99") if jc else None),
                fmt_us(rc.get("p999") if rc else None), fmt_us(jc.get("p999") if jc else None),
                "yes" if sat else "no"])
        add_table(doc, ["Operation", "target ops/s", "Rust p99", "Java p99",
                        "Rust p99.9", "Java p99.9", "saturated"], orows,
                  widths=[1.8, 1.1, 0.9, 0.9, 0.95, 0.95, 0.95])

    # ---- validity & caveats ----
    doc.add_heading("Validity & caveats", level=1)
    contaminated = [r["cell_id"] + "/" + r["client"] for r in recs if r.get("_env", {}).get("contamination")]
    low_tail = sum(1 for cid, e in A.items()
                   if any(c.get("samples_above_p999_min", 0) < 1000 for c in e["clients"].values()))
    for txt in [
        "Total errors across all runs: %d." % total_err,
        "Contamination flags (member migration / slow-op / Full-GC during a window): %d." % len(contaminated),
        "Cells with <1000 samples above p99.9 (tail estimate weaker, typically the C=1 latency-floor cells): %d."
        % low_tail,
        "Co-located rig: cluster and client share one host with disjoint core pinning (members %s, client %s). "
        "They still share L3 cache, memory bandwidth and the loopback NIC, so absolute throughput is higher "
        "than a two-machine LAN deployment — but both clients see identical conditions, so the comparison is fair."
        % (pin.get("members"), pin.get("client")),
        "CPU governor is hypervisor-managed on the virtualized guest; achieved frequency is recorded in provenance.",
        "Statistical depth: median across forks×trials with a 95% bootstrap CI (n is modest for a fresh broad run; "
        "wider CIs are reported honestly rather than hidden).",
        "Java RSS includes the whole JVM (heap + metaspace + code cache + stacks); the Rust figure is the whole "
        "process RSS. Both are apples-to-apples “memory the client process needs”.",
    ]:
        doc.add_paragraph(txt, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Appendix — full provenance", level=2)
    import json as _json
    code = doc.add_paragraph(_json.dumps(prov, indent=2))
    for run in code.runs:
        run.font.name = "Consolas"; run.font.size = Pt(8)

    doc.save(out_path)
    print("wrote", out_path, "(%d cells, %d runs)" % (len(A), len(recs)))
    return True


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_dir", required=True)
    ap.add_argument("--out", default=None)
    ap.add_argument("--charts", default=None)
    args = ap.parse_args()
    out = args.out or os.path.join(args.in_dir, "HazelRust_vs_Java_H2H_Report.docx")
    charts = args.charts or os.path.join(args.in_dir, "charts")
    build_docx(args.in_dir, out, charts)


if __name__ == "__main__":
    main()
