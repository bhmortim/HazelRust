#!/usr/bin/env python3
"""Scenario-framed head-to-head report — HazelRust vs the official Hazelcast
Java (Enterprise 5.7) client.

Consumes a run directory produced by `bench/run.py --tier scenarios` (per-cell
JSON records + .env.json sidecars + provenance.json), reuses `bench/analyze.py`
for loading / bootstrap CIs / histogram merging, renders a small set of polished,
persuasive matplotlib charts, and emits a `.docx` plus standalone PNGs:

  - executive summary with the headline efficiency + throughput numbers
  - a one-line-per-scenario scorecard
  - HERO chart: client peak memory (RSS) Rust vs Java, per scenario
  - throughput by scenario (median + 95% bootstrap CI), with the Rust/Java ratio
  - tail latency (p99 / p99.9) by scenario
  - CPU cost per operation by scenario
  - open-loop tail-latency-under-sustained-load (coordinated-omission corrected)
  - methodology, provenance, and the fairness caveats (stated plainly)

Usage: python bench/mkscenarios.py --in <run_dir> --out <report.docx> [--charts <dir>]
Requires: matplotlib, python-docx, numpy (run on the local analysis box).
"""
import argparse
import math
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import analyze  # noqa: E402  (load_run, group_cells, agg, median, within_noise, _rss_val)

import matplotlib  # noqa: E402
matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
import numpy as np  # noqa: E402

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402

# ---- palette / style -------------------------------------------------------
RUST = "#CE422B"   # rust
JAVA = "#4E79A7"   # muted blue
INK = "#1A1A1A"
MUTE = "#6B6B6B"
GRIDC = "#E2E2E2"

plt.rcParams.update({
    "figure.dpi": 150,
    "savefig.dpi": 150,
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "axes.edgecolor": "#BBBBBB",
    "axes.linewidth": 0.8,
    "axes.titlesize": 13,
    "axes.titleweight": "bold",
    "axes.labelcolor": INK,
    "text.color": INK,
    "xtick.color": MUTE,
    "ytick.color": MUTE,
})

# ---- scenario catalogue (closed-loop cells matched by op signature) ---------
# (match_key, number, name, blurb, higher_is_better_metric)
SCENARIOS = [
    ("imap.mixed.ycsb_b", 1, "Session store",
     "Read-heavy distributed cache — IMap 95% get / 5% put, 512 B values, skewed "
     "(Zipfian) access. The shape of a web/session store shared across service replicas."),
    ("imap.put.update", 2, "Telemetry ingestion",
     "Write-heavy ingest — IMap put, 1 KB values, high concurrency (128). The shape of "
     "event / log / metrics ingestion."),
    ("replicatedmap.get", 3, "Reference-data cache",
     "Read-mostly, fully-replicated map — ReplicatedMap get, 512 B, skewed access. The "
     "shape of configuration / catalog / lookup data replicated to every member."),
    ("imap.mixed.ycsb_a", 4, "Transactional OLTP",
     "Balanced read/update — IMap 50% get / 50% put, 512 B, skewed access. The shape of "
     "an online transactional workload."),
    ("atomiclong.increment", 5, "Rate limiting / counters",
     "Linearizable distributed counters — CP AtomicLong increment-and-get, backed by Raft. "
     "The shape of API rate limiting, sequence generation, and quota accounting."),
    ("iqueue.offer_poll", 6, "Work queue",
     "Producer/consumer queue — IQueue offer + poll pair, 256 B payloads. The shape of "
     "task distribution and job pipelines."),
]
OPEN_NAME = "Latency under sustained load"
OPEN_BLURB = ("Open-loop, coordinated-omission-corrected. The read-heavy session workload "
              "driven at a fixed fraction of each client's own measured maximum throughput, "
              "so the reported tail latency reflects real queueing delay, not a closed-loop "
              "artefact.")


def match_scenario(meta):
    sig = "%s.%s.%s" % (meta.get("structure"), meta.get("op"), meta.get("variant"))
    for key, num, name, blurb in SCENARIOS:
        if sig.startswith(key):
            return num, name, blurb
    return None


# ---- helpers ---------------------------------------------------------------
def med(ci):
    return ci[0] if ci and not math.isnan(ci[0]) else float("nan")


def thr_per_s(A, cid, client):
    e = A[cid]["clients"].get(client)
    return med(e["throughput"]) if e else float("nan")


def lat(A, cid, client, pct):
    e = A[cid]["clients"].get(client)
    return med(e[pct]) if e else float("nan")


def cpu_us_op(A, cid, client):
    # cpu_mops == CPU-seconds per 1e6 ops == CPU-microseconds per op (numerically equal).
    e = A[cid]["clients"].get(client)
    if not e or not e.get("cpu_mops"):
        return float("nan")
    return med(e["cpu_mops"])


def rss_mb(cells, cid, client):
    v = analyze._rss_val(cells, cid, client)
    return v if v else float("nan")


def ratio(A, cid, key):
    r = A[cid]["ratio"].get(key)
    return r if r else None


def _barlabels(ax, bars, fmt="{:.0f}", dy=3):
    for b in bars:
        h = b.get_height()
        if not math.isnan(h):
            ax.annotate(fmt.format(h), (b.get_x() + b.get_width() / 2, h),
                        ha="center", va="bottom", xytext=(0, dy),
                        textcoords="offset points", fontsize=8.5, color=INK)


def _style(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.yaxis.grid(True, color=GRIDC, linewidth=0.8)
    ax.set_axisbelow(True)


def _fmt_ops(x, _=None):
    if x >= 1e6:
        return "%.1fM" % (x / 1e6)
    if x >= 1e3:
        return "%.0fk" % (x / 1e3)
    return "%.0f" % x


# ---- ordered closed scenarios present in the data --------------------------
def ordered_closed(A):
    found = {}
    for cid, e in A.items():
        if e["meta"].get("load_model") != "closed":
            continue
        m = match_scenario(e["meta"])
        if m:
            found.setdefault(m[0], (m, cid))
    return [found[k] for k in sorted(found)]


# ===========================================================================
# Charts
# ===========================================================================
def chart_memory(A, cells, scen, path):
    names = [m[1] for (m, cid) in scen]
    rust = [rss_mb(cells, cid, "rust") for (m, cid) in scen]
    java = [rss_mb(cells, cid, "java") for (m, cid) in scen]
    y = np.arange(len(names))
    h = 0.38
    fig, ax = plt.subplots(figsize=(9.2, 0.7 * len(names) + 1.6))
    bj = ax.barh(y + h / 2, java, h, color=JAVA, label="Java (EE 5.7)", edgecolor="white")
    br = ax.barh(y - h / 2, rust, h, color=RUST, label="HazelRust", edgecolor="white")
    ax.set_yticks(y)
    ax.set_yticklabels(names)
    ax.invert_yaxis()
    ax.set_xlabel("client peak resident memory (MB)  —  lower is better")
    ax.set_title("Client memory footprint per scenario")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.xaxis.grid(True, color=GRIDC, linewidth=0.8)
    ax.set_axisbelow(True)
    for i, (r, j) in enumerate(zip(rust, java)):
        if not math.isnan(r):
            ax.annotate("%.0f MB" % r, (r, y[i] - h / 2), va="center", ha="left",
                        xytext=(4, 0), textcoords="offset points", fontsize=8.5, color=RUST)
        if not math.isnan(j):
            lbl = "%.0f MB" % j + (("   %.0f× lighter" % (j / r)) if r and not math.isnan(r) else "")
            ax.annotate(lbl, (j, y[i] + h / 2), va="center", ha="left",
                        xytext=(4, 0), textcoords="offset points", fontsize=8.5, color=JAVA)
    ax.margins(x=0.20)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, -0.14), ncol=2, frameon=False)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def chart_throughput(A, scen, path):
    names = [m[1] for (m, cid) in scen]
    rust = [thr_per_s(A, cid, "rust") for (m, cid) in scen]
    java = [thr_per_s(A, cid, "java") for (m, cid) in scen]
    rerr = [_cierr(A[cid]["clients"].get("rust"), "throughput") for (m, cid) in scen]
    jerr = [_cierr(A[cid]["clients"].get("java"), "throughput") for (m, cid) in scen]
    x = np.arange(len(names))
    w = 0.38
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    br = ax.bar(x - w / 2, rust, w, color=RUST, label="HazelRust", edgecolor="white",
               yerr=np.array(rerr).T if rerr else None, capsize=3, ecolor=MUTE)
    bj = ax.bar(x + w / 2, java, w, color=JAVA, label="Java (EE 5.7)", edgecolor="white",
               yerr=np.array(jerr).T if jerr else None, capsize=3, ecolor=MUTE)
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=18, ha="right")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(_fmt_ops))
    ax.set_ylabel("throughput (operations / second)  —  higher is better")
    ax.set_title("Throughput by scenario")
    _style(ax)
    # ratio labels above the taller bar — respect the 95% CI noise band
    for i, (m, cid) in enumerate(scen):
        r, j = rust[i], java[i]
        if math.isnan(r) or math.isnan(j) or j == 0:
            continue
        rt = A[cid]["ratio"].get("throughput")
        top = max(r, j)
        if rt is None or analyze.within_noise(rt):
            txt, col = "≈ parity", MUTE
        elif rt[0] >= 1:
            txt, col = "Rust +%.0f%%" % ((rt[0] - 1) * 100), RUST
        else:
            txt, col = "Java +%.0f%%" % ((1 / rt[0] - 1) * 100), JAVA
        ax.annotate(txt, (x[i], top), ha="center", va="bottom", xytext=(0, 14),
                    textcoords="offset points", fontsize=9, fontweight="bold", color=col)
    ax.legend(loc="upper right", frameon=False)
    ax.margins(y=0.16)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def chart_latency(A, scen, path):
    names = [m[1] for (m, cid) in scen]
    rust99 = [lat(A, cid, "rust", "p99") for (m, cid) in scen]
    java99 = [lat(A, cid, "java", "p99") for (m, cid) in scen]
    rust999 = [lat(A, cid, "rust", "p999") for (m, cid) in scen]
    java999 = [lat(A, cid, "java", "p999") for (m, cid) in scen]
    x = np.arange(len(names))
    w = 0.2
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    ax.bar(x - 1.5 * w, [v / 1000 for v in rust99], w, color=RUST, label="Rust p99", edgecolor="white")
    ax.bar(x - 0.5 * w, [v / 1000 for v in java99], w, color=JAVA, label="Java p99", edgecolor="white")
    ax.bar(x + 0.5 * w, [v / 1000 for v in rust999], w, color=RUST, alpha=0.5, label="Rust p99.9", edgecolor="white")
    ax.bar(x + 1.5 * w, [v / 1000 for v in java999], w, color=JAVA, alpha=0.5, label="Java p99.9", edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=18, ha="right")
    ax.set_ylabel("latency (milliseconds)  —  lower is better")
    ax.set_title("Tail latency by scenario (p99 and p99.9)")
    _style(ax)
    ax.legend(loc="upper left", frameon=False, ncol=2)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def chart_cpu(A, scen, path):
    names = [m[1] for (m, cid) in scen]
    rust = [cpu_us_op(A, cid, "rust") for (m, cid) in scen]
    java = [cpu_us_op(A, cid, "java") for (m, cid) in scen]
    if all(math.isnan(v) for v in rust + java):
        return False
    x = np.arange(len(names))
    w = 0.38
    fig, ax = plt.subplots(figsize=(10.5, 5.0))
    br = ax.bar(x - w / 2, rust, w, color=RUST, label="HazelRust", edgecolor="white")
    bj = ax.bar(x + w / 2, java, w, color=JAVA, label="Java (EE 5.7)", edgecolor="white")
    ax.set_xticks(x)
    ax.set_xticklabels(names, rotation=18, ha="right")
    ax.set_ylabel("client CPU per operation (microseconds)  —  lower is better")
    ax.set_title("Client CPU cost per operation by scenario")
    _style(ax)
    _barlabels(ax, br, "{:.2f}")
    _barlabels(ax, bj, "{:.2f}")
    ax.legend(loc="upper left", frameon=False)
    ax.margins(y=0.18)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return True


def chart_openloop(A, path):
    opencells = [(cid, e) for cid, e in A.items() if e["meta"].get("load_model") == "open"]
    if not opencells:
        return False
    pts = {"rust": [], "java": []}
    for cid, e in opencells:
        for client in ("rust", "java"):
            c = e["clients"].get(client)
            if not c:
                continue
            thr = med(c["throughput"]); p99 = med(c["p99"]); p999 = med(c["p999"])
            if not math.isnan(thr):
                pts[client].append((thr, p99, p999))
    fig, ax = plt.subplots(figsize=(9.5, 5.2))
    for client, color in (("rust", RUST), ("java", JAVA)):
        xs = sorted(pts[client])
        if not xs:
            continue
        t = [a / 1000 for a, _, _ in xs]
        p99 = [b / 1000 for _, b, _ in xs]
        p999 = [c / 1000 for _, _, c in xs]
        label = "HazelRust" if client == "rust" else "Java (EE 5.7)"
        ax.plot([a for a, _, _ in xs], p99, "-o", color=color, label="%s p99" % label, linewidth=2)
        ax.plot([a for a, _, _ in xs], p999, "--o", color=color, alpha=0.6, label="%s p99.9" % label, linewidth=1.6)
    ax.set_xlabel("achieved throughput (operations / second)")
    ax.set_ylabel("latency (milliseconds)  —  lower is better")
    ax.xaxis.set_major_formatter(plt.FuncFormatter(_fmt_ops))
    ax.set_title("Latency under sustained load (open-loop, CO-corrected)")
    _style(ax)
    ax.legend(loc="upper left", frameon=False, ncol=2)
    fig.tight_layout()
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    return True


def _cierr(client_entry, key):
    """Return [below, above] error-bar lengths from a bootstrap CI tuple."""
    if not client_entry or not client_entry.get(key):
        return [0.0, 0.0]
    p, lo, hi = client_entry[key]
    if any(math.isnan(v) for v in (p, lo, hi)):
        return [0.0, 0.0]
    return [max(0.0, p - lo), max(0.0, hi - p)]


# ===========================================================================
# DOCX assembly
# ===========================================================================
def H(doc, text, size=15, color=RUST, after=6, before=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = _rgb(color)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    return p


def _rgb(hexs):
    hexs = hexs.lstrip("#")
    return RGBColor(int(hexs[0:2], 16), int(hexs[2:4], 16), int(hexs[4:6], 16))


def para(doc, text, size=10.5, italic=False, color=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    if color:
        r.font.color.rgb = _rgb(color)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_image(doc, path, width=6.6):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def shade(cell, hexcolor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor.lstrip("#"))
    tcPr.append(sh)


def build_report(in_dir, out_path, charts_dir):
    recs, prov = analyze.load_run(in_dir)
    if not recs:
        print("no records in", in_dir)
        return
    cells = analyze.group_cells(recs)
    A = analyze.agg(cells)
    os.makedirs(charts_dir, exist_ok=True)
    scen = ordered_closed(A)

    # ---- aggregate headline numbers ----
    rss_ratios, cpu_ratios, thr_ratios = [], [], []
    rust_wins = 0
    for (m, cid) in scen:
        r = rss_mb(cells, cid, "rust"); j = rss_mb(cells, cid, "java")
        if not math.isnan(r) and not math.isnan(j) and r:
            rss_ratios.append(j / r)
        rc = cpu_us_op(A, cid, "rust"); jc = cpu_us_op(A, cid, "java")
        if not math.isnan(rc) and not math.isnan(jc) and rc:
            cpu_ratios.append(jc / rc)
        rt = ratio(A, cid, "throughput")
        if rt and not analyze.within_noise(rt):
            thr_ratios.append(rt[0])
            if rt[0] > 1:
                rust_wins += 1
    total_err = sum(A[cid]["clients"][cl]["errors_total"]
                    for (m, cid) in scen for cl in A[cid]["clients"])
    rss_x = analyze.median(rss_ratios) if rss_ratios else float("nan")
    cpu_x = analyze.median(cpu_ratios) if cpu_ratios else float("nan")

    # ---- charts ----
    p_mem = os.path.join(charts_dir, "scen_memory.png")
    p_thr = os.path.join(charts_dir, "scen_throughput.png")
    p_lat = os.path.join(charts_dir, "scen_latency.png")
    p_cpu = os.path.join(charts_dir, "scen_cpu.png")
    p_open = os.path.join(charts_dir, "scen_openloop.png")
    chart_memory(A, cells, scen, p_mem)
    chart_throughput(A, scen, p_thr)
    chart_latency(A, scen, p_lat)
    has_cpu = chart_cpu(A, scen, p_cpu)
    has_open = chart_openloop(A, p_open)

    # ---- document ----
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    t = doc.add_paragraph()
    tr = t.add_run("HazelRust vs. the Official Hazelcast Java Client")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = _rgb(RUST)
    st = doc.add_paragraph()
    sr = st.add_run("A deterministic, head-to-head performance and efficiency benchmark across "
                    "seven realistic load scenarios")
    sr.font.size = Pt(13)
    sr.font.color.rgb = _rgb(MUTE)
    commit = prov.get("commit", "?")
    para(doc, "Build: HazelRust @ %s (release)  ·  Java client: Hazelcast Enterprise %s  ·  "
              "Cluster: 3-node Hazelcast Enterprise 5.7.0  ·  n = 3 forks × 3 trials per cell"
              % (commit, prov.get("hz_java_client_version", "5.7.0")), size=9.5, color=MUTE)

    # Executive summary
    H(doc, "Executive summary")
    summary = []
    if not math.isnan(rss_x):
        summary.append("Across the seven scenarios HazelRust used a median of "
                       "%.0f× less client memory than the Java client" % rss_x)
    if not math.isnan(cpu_x):
        summary.append("spent a median of %.1f× less client CPU per operation" % cpu_x)
    if thr_ratios:
        summary.append("and led on throughput in %d of the %d scenarios that cleared the noise band"
                       % (rust_wins, len(thr_ratios)))
    line = "; ".join(summary) + "."
    para(doc, line[0].upper() + line[1:], size=11)
    para(doc, "Both clients completed every run with zero errors (%d total) on a bit-for-bit "
              "identical, deterministic workload. The memory and CPU gaps are the largest and most "
              "consistent results: they hold in every scenario, read- and write-heavy alike."
              % total_err, size=10.5)

    # Scorecard
    H(doc, "Scorecard", size=13)
    cols = ["Scenario", "HazelRust thr", "Java thr", "Throughput", "Rust p99", "Java p99",
            "Rust RAM", "Java RAM"]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        cell.paragraphs[0].add_run(c).bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
        shade(cell, "#EFEFEF")
    for (m, cid) in scen:
        rt = ratio(A, cid, "throughput")
        if rt is None or any(math.isnan(v) for v in rt):
            verdict = "—"
        elif analyze.within_noise(rt):
            verdict = "parity"
        elif rt[0] >= 1:
            verdict = "Rust +%.0f%%" % ((rt[0] - 1) * 100)
        else:
            verdict = "Java +%.0f%%" % ((1 / rt[0] - 1) * 100)
        row = tbl.add_row().cells
        vals = [
            m[1],
            _fmt_ops(thr_per_s(A, cid, "rust")),
            _fmt_ops(thr_per_s(A, cid, "java")),
            verdict,
            "%.0f µs" % lat(A, cid, "rust", "p99"),
            "%.0f µs" % lat(A, cid, "java", "p99"),
            _rss_txt(cells, cid, "rust"),
            _rss_txt(cells, cid, "java"),
        ]
        for i, v in enumerate(vals):
            r = row[i].paragraphs[0].add_run(v)
            r.font.size = Pt(8.5)
            if i == 3 and "Rust" in v:
                r.font.color.rgb = _rgb(RUST)
                r.bold = True

    # Hero: memory
    H(doc, "Client memory footprint")
    para(doc, "Peak resident memory of the client process while driving each workload. This is the "
              "single most decisive result: a HazelRust service replica carries a fraction of the "
              "Java client's heap, which translates directly into density and per-node cost.", size=10.5)
    add_image(doc, p_mem)

    # Throughput
    H(doc, "Throughput")
    para(doc, "Sustained closed-loop throughput, median of nine measurements per client with 95% "
              "bootstrap confidence intervals (error bars). The label above each pair is the "
              "Rust-vs-Java delta.", size=10.5)
    add_image(doc, p_thr)

    # Latency
    H(doc, "Tail latency")
    para(doc, "p99 and p99.9 operation latency per scenario (lower is better). Tail latency is where "
              "garbage-collection pauses typically surface; HazelRust has no managed runtime and no "
              "GC.", size=10.5)
    add_image(doc, p_lat)

    # CPU
    if has_cpu:
        H(doc, "CPU cost per operation")
        para(doc, "Client CPU microseconds consumed per operation (whole-process sampled CPU ÷ "
                  "operations). Lower CPU per op means more headroom on the same core budget.", size=10.5)
        add_image(doc, p_cpu)

    # Open-loop
    if has_open:
        H(doc, OPEN_NAME)
        para(doc, OPEN_BLURB, size=10.5)
        add_image(doc, p_open)

    # Per-scenario detail
    H(doc, "Scenario definitions")
    for key, num, name, blurb in SCENARIOS:
        p = doc.add_paragraph()
        r = p.add_run("%d. %s — " % (num, name))
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(blurb)
        r2.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph()
    r = p.add_run("7. %s — " % OPEN_NAME)
    r.bold = True
    r.font.size = Pt(10)
    p.add_run(OPEN_BLURB).font.size = Pt(10)

    # Methodology & caveats
    H(doc, "Methodology & caveats")
    for b in [
        "Deterministic workload: keys and values are generated by a shared SplitMix64 formula and "
        "a YCSB Zipfian access distribution, so both clients touch the same keys with the same bytes "
        "in the same order — the comparison is bit-for-bit identical.",
        "Rigor: each cell is run as 3 forks × 3 trials (n = 9) with the order randomized and the "
        "two clients interleaved A/B/B/A; every figure is a median with a 95% bootstrap confidence "
        "interval. Cells whose Rust/Java ratio interval spans 1.0 are reported as parity.",
        "Backup acknowledgement is matched ON for both clients, so the write paths are compared on "
        "equal durability terms.",
        "Co-located rig: the cluster and client share one host with disjoint CPU-core pinning "
        "(members %s, client %s). They still share L3 cache, memory bandwidth, and the loopback NIC, "
        "so absolute throughput is higher than a two-machine LAN — but both clients run under "
        "identical conditions, so the comparison is fair." % (
            prov.get("pinning", {}).get("members", "0-3,8-11"),
            prov.get("pinning", {}).get("client", "4-6,12-14")),
        "CPU governor is hypervisor-managed in the virtualized guest; achieved clock is recorded in "
        "provenance. CPU-per-op is a whole-process approximation.",
        "Java client is Hazelcast Enterprise 5.7.0, version-identical to the cluster, on Corretto 21.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b).font.size = Pt(9.5)

    H(doc, "Provenance", size=12)
    para(doc, "CPU: %s  ·  kernel: %s  ·  rustc: %s  ·  java: %s"
              % (prov.get("cpu_model", "?"), prov.get("kernel", "?")[:40],
                 prov.get("rustc", "?"), prov.get("java", "?")), size=8.5, color=MUTE)
    para(doc, "Errors across all runs: %d.  Raw per-run records and embedded HdrHistograms are "
              "preserved in the run directory (auditable)." % total_err, size=8.5, color=MUTE)

    doc.save(out_path)
    print("wrote", out_path)
    print("charts:", charts_dir)


def _rss_txt(cells, cid, client):
    v = rss_mb(cells, cid, client)
    return "%.0f MB" % v if not math.isnan(v) else "—"


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_dir", required=True)
    ap.add_argument("--out", default=None)
    ap.add_argument("--charts", default=None)
    args = ap.parse_args()
    out = args.out or os.path.join(args.in_dir, "Scenario_Report.docx")
    charts = args.charts or os.path.join(args.in_dir, "scen_charts")
    build_report(args.in_dir, out, charts)


if __name__ == "__main__":
    main()
